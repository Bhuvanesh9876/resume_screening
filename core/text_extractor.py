"""
Text Extractor Module

This module handles the extraction of text from PDF and DOCX files.
It employs multiple strategies (PyMuPDF, pdfplumber, pdfminer) to ensure
robust text extraction from various PDF formats.
"""

import io
import re
from typing import Any
import fitz
import pdfplumber
import chardet
from pdfminer.high_level import extract_text as pdfminer_extract
try:
    from PIL import Image
    import pytesseract
    from pdf2image import convert_from_bytes
except ImportError:
    Image = None
    pytesseract = None
    convert_from_bytes = None

def _extract_with_pymupdf(file: Any) -> str:
    """Extract text using PyMuPDF (fitz)."""
    try:
        if hasattr(file, "read"):
            data = file.read()
            file.seek(0)
        else:
            with open(file, "rb") as f:
                data = f.read()

        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                pages.append(text.strip())
        doc.close()
        return "\n".join(pages)
    except Exception:
        return ""

def _extract_with_pdfplumber(file: Any) -> str:
    """Extract text using pdfplumber."""
    try:
        if hasattr(file, "seek"):
            file.seek(0)
        with pdfplumber.open(file) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    pages.append(text.strip())
        return "\n".join(pages)
    except Exception:
        return ""

def _extract_with_pdfminer(file: Any) -> str:
    """Extract text using pdfminer.six."""
    try:
        if hasattr(file, "read"):
            file.seek(0)
            data = file.read()
            file.seek(0)
            return pdfminer_extract(io.BytesIO(data))
        return pdfminer_extract(file)
    except Exception:
        return ""

def _extract_from_docx(file: Any) -> str:
    """Extract text from DOCX files using python-docx.

    Reads both paragraphs AND tables because many resumes use table
    layouts for education, experience, and contact sections.
    """
    try:
        from docx import Document
        if hasattr(file, "read"):
            file.seek(0)
            doc = Document(file)
        else:
            doc = Document(file)

        parts: list[str] = []

        # 1. Paragraphs (normal body text)
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        # 2. Tables (common in formatted resumes — education, skills, etc.)
        for table in doc.tables:
            for row in table.rows:
                cells = []
                seen: set[str] = set()
                for cell in row.cells:
                    ct = cell.text.strip()
                    # Deduplicate merged cells that report the same text
                    if ct and ct not in seen:
                        seen.add(ct)
                        cells.append(ct)
                if cells:
                    parts.append("  ".join(cells))

        return "\n".join(parts)
    except Exception:
        return ""

def _extract_from_txt(file: Any) -> str:
    """Extract text from TXT files with encoding detection."""
    try:
        if hasattr(file, "read"):
            raw_data = file.read()
            file.seek(0)
        else:
            with open(file, "rb") as f:
                raw_data = f.read()
        
        result = chardet.detect(raw_data)
        encoding = result['encoding'] or 'utf-8'
        return raw_data.decode(encoding)
    except Exception:
        return ""

def _extract_with_ocr(file: Any) -> str:
    """Extract text from images using OCR (pytesseract)."""
    if not pytesseract or not Image:
        return ""
    try:
        if hasattr(file, "read"):
            file.seek(0)
            data = file.read()
            file.seek(0)
            img = Image.open(io.BytesIO(data))
        else:
            img = Image.open(file)
        
        text = pytesseract.image_to_string(img)
        return text
    except Exception:
        return ""

def _extract_scanned_pdf(file: Any) -> str:
    """Fallback for scanned PDFs: convert to images and then OCR."""
    if not convert_from_bytes or not pytesseract:
        return ""
    try:
        if hasattr(file, "read"):
            file.seek(0)
            data = file.read()
            file.seek(0)
        else:
            with open(file, "rb") as f:
                data = f.read()
        
        images = convert_from_bytes(data)
        full_text = []
        for img in images:
            text = pytesseract.image_to_string(img)
            full_text.append(text)
        return "\n".join(full_text)
    except Exception:
        return ""

def _post_clean(text: str) -> str:
    """
    Clean and normalize extracted text.
    Removes invisible characters, weird hyphenation, and excessive newlines.
    """
    if not text:
        return ""
    text = text.replace("\xa0", " ").replace("\u200b", "")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.replace("\x00", "")
    return text.strip()

def extract_text(file: Any) -> str:
    """
    Main entry point for text extraction.
    Handles PDF, DOCX, TXT, and Images using multiple strategies.
    """
    filename = ""
    if hasattr(file, "name"):
        filename = file.name.lower()

    # Handle DOCX
    if filename.endswith(".docx"):
        text = _extract_from_docx(file)
        return _post_clean(text)
    
    # Handle TXT
    if filename.endswith(".txt"):
        text = _extract_from_txt(file)
        return _post_clean(text)

    # Handle Images
    if filename.endswith((".png", ".jpg", ".jpeg")):
        text = _extract_with_ocr(file)
        return _post_clean(text)

    # Handle PDF with tiered extraction + OCR Fallback
    # 1. PyMuPDF (Fast)
    text = _extract_with_pymupdf(file)
    if len(text.strip()) > 100:
        return _post_clean(text)

    # 2. pdfplumber (Better layout)
    text = _extract_with_pdfplumber(file)
    if len(text.strip()) > 100:
        return _post_clean(text)

    # 3. pdfminer (Fallback)
    text = _extract_with_pdfminer(file)
    if len(text.strip()) > 100:
        return _post_clean(text)
    
    # 4. OCR Fallback for Scanned PDFs
    text = _extract_scanned_pdf(file)
    return _post_clean(text)
